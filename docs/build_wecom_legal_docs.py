#!/usr/bin/env python3
"""Generate upload-ready WeCom legal documents for Kellai.

The generated DOCX files are editable source documents. PDF is emitted by the
rendering workflow so the uploaded copy remains visually stable.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence
from xml.sax.saxutils import escape

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "docs" / "export"
PDF_OUT = ROOT / "output" / "pdf"

PROVIDER = "成都修茈科技有限公司"
PRODUCT = "客来来"
APP_ID = "ww4ede93d0329bf40d"
VERSION = "V1.0"
DATE_CN = "2026年7月14日"
WEBSITE = "https://xiu-ci.com"
EMAIL = "support@xiu-ci.com"

NAVY = "163A5F"
BLUE = "2E6F9E"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D8DEE5"
DARK = "20262E"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    repeat_table_header(row)
    prevent_row_split(row)


def set_font(run, *, name="Arial Unicode MS", east_asia="Arial Unicode MS", size=10.5, bold=None, color=DARK) -> None:
    # LibreOffice's DOCX renderer does not consistently honor only w:eastAsia.
    # Use a CJK-capable family for ascii/hAnsi too so uploaded PDFs never show tofu.
    run.font.name = east_asia
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, *, name="Arial Unicode MS", east_asia="Arial Unicode MS", size=10.5, bold=False, color=DARK) -> None:
    style.font.name = east_asia
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=8.5, color="68727D")


def add_numbering_definition(doc: Document, num_id: int = 40) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = num_id

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "（%1）")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "420")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    doc._legal_num_counter = num_id
    return abstract_id


def create_numbering_instance(doc: Document, abstract_id: int) -> int:
    doc._legal_num_counter += 1
    num_id = doc._legal_num_counter
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_elem])
    p_pr.append(num_pr)


def setup_document(title: str, short_title: str) -> tuple[Document, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=10.5)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.widow_control = True

    for style_name, east_asia, size, bold, color in (
        ("Title", "Hiragino Sans GB", 20, True, NAVY),
        ("Heading 1", "Hiragino Sans GB", 13, True, NAVY),
        ("Heading 2", "Hiragino Sans GB", 11, True, BLUE),
    ):
        style = styles[style_name]
        set_style_font(style, east_asia=east_asia, size=size, bold=bold, color=color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    styles["Heading 1"].paragraph_format.space_before = Pt(11)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.space_before = Pt(7)
    styles["Heading 2"].paragraph_format.space_after = Pt(3)

    if "Legal Metadata" not in styles:
        meta = styles.add_style("Legal Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Legal Metadata"]
    set_style_font(meta, size=9.5, color="4D5965")
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Legal Lead" not in styles:
        lead = styles.add_style("Legal Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Legal Lead"]
    set_style_font(lead, size=10.5, color=DARK)
    lead.paragraph_format.left_indent = Inches(0.18)
    lead.paragraph_format.right_indent = Inches(0.18)
    lead.paragraph_format.space_before = Pt(8)
    lead.paragraph_format.space_after = Pt(9)
    lead.paragraph_format.line_spacing = 1.3
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Legal Notice" not in styles:
        notice = styles.add_style("Legal Notice", WD_STYLE_TYPE.PARAGRAPH)
    else:
        notice = styles["Legal Notice"]
    set_style_font(notice, size=9.5, color=NAVY)
    notice.paragraph_format.left_indent = Inches(0.16)
    notice.paragraph_format.right_indent = Inches(0.16)
    notice.paragraph_format.space_before = Pt(5)
    notice.paragraph_format.space_after = Pt(7)
    notice.paragraph_format.line_spacing = 1.2

    if "Legal Numbered" not in styles:
        numbered = styles.add_style("Legal Numbered", WD_STYLE_TYPE.PARAGRAPH)
    else:
        numbered = styles["Legal Numbered"]
    set_style_font(numbered, size=10.5)
    numbered.paragraph_format.space_after = Pt(4)
    numbered.paragraph_format.line_spacing = 1.22
    numbered.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{PRODUCT}  |  {PROVIDER}")
    set_font(r, east_asia="Hiragino Sans GB", size=8.5, bold=True, color=NAVY)
    r2 = p.add_run(f"    {short_title}")
    set_font(r2, size=8.5, color="68727D")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("第 ")
    set_font(fr, size=8.5, color="68727D")
    set_page_field(fp, "PAGE")
    fr = fp.add_run(" 页 / 共 ")
    set_font(fr, size=8.5, color="68727D")
    set_page_field(fp, "NUMPAGES")
    fr = fp.add_run(" 页")
    set_font(fr, size=8.5, color="68727D")

    props = doc.core_properties
    props.title = title
    props.subject = f"{PRODUCT}企业微信第三方应用法律文件"
    props.author = PROVIDER
    props.company = PROVIDER
    props.comments = "Generated for WeCom service provider upload."

    num_id = add_numbering_definition(doc)
    return doc, num_id


def add_title_block(doc: Document, title: str, description: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.add_run(title)

    for line in (
        f"服务提供者：{PROVIDER}",
        f"适用产品：{PRODUCT}（企业微信第三方应用标识：{APP_ID}）",
        f"版本：{VERSION}    更新日期：{DATE_CN}    生效日期：{DATE_CN}",
    ):
        doc.add_paragraph(line, style="Legal Metadata")

    p = doc.add_paragraph(style="Legal Lead")
    set_paragraph_fill(p, LIGHT_BLUE)
    r = p.add_run(description)
    set_font(r, size=10.5, color=DARK)


def set_paragraph_fill(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold_prefix: str = "", style=None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)


def add_numbered(doc: Document, abstract_id: int, items: Iterable[str]) -> None:
    num_id = create_numbering_instance(doc, abstract_id)
    for item in items:
        p = doc.add_paragraph(style="Legal Numbered")
        apply_num(p, num_id)
        r = p.add_run(item)
        set_font(r)


def add_notice(doc: Document, text: str, *, label="重要提示：") -> None:
    p = doc.add_paragraph(style="Legal Notice")
    set_paragraph_fill(p, LIGHT_GRAY)
    r = p.add_run(label)
    set_font(r, east_asia="Hiragino Sans GB", size=9.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, size=9.5, color=NAVY)


def add_info_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, east_asia="Hiragino Sans GB", size=9, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            if row_idx % 2:
                set_cell_shading(cell, "FAFBFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_font(r, size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def save(doc: Document, filename: str) -> Path:
    DOCX_OUT.mkdir(parents=True, exist_ok=True)
    path = DOCX_OUT / filename
    doc.save(path)
    return path


def iter_docx_blocks(parent):
    """Yield top-level paragraphs and tables in document order."""
    if not isinstance(parent, _Document):
        raise TypeError("PDF export currently expects a Document")
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, parent)


def register_pdf_fonts() -> tuple[str, str]:
    body_name = "KellaiArialUnicode"
    heading_name = "KellaiHeiti"
    if body_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(
            TTFont(body_name, "/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
        )
    if heading_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(
            TTFont(heading_name, "/System/Library/Fonts/STHeiti Medium.ttc", subfontIndex=0)
        )
    pdfmetrics.registerFontFamily(
        body_name,
        normal=body_name,
        bold=heading_name,
        italic=body_name,
        boldItalic=heading_name,
    )
    return body_name, heading_name


def _docx_num_id(paragraph: DocxParagraph) -> int | None:
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
        return None
    return int(p_pr.numPr.numId.val)


def export_pdf_from_docx(docx_path: Path, pdf_path: Path, short_title: str) -> Path:
    body_font, heading_font = register_pdf_fonts()
    source = Document(docx_path)
    PDF_OUT.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "KellaiBody",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=9.6,
        leading=14.2,
        textColor=colors.HexColor(f"#{DARK}"),
        alignment=TA_JUSTIFY,
        wordWrap="CJK",
        spaceAfter=5.2,
        allowWidows=0,
        allowOrphans=0,
    )
    title_style = ParagraphStyle(
        "KellaiTitle",
        parent=body,
        fontName=heading_font,
        fontSize=19,
        leading=25,
        textColor=colors.HexColor(f"#{NAVY}"),
        alignment=TA_CENTER,
        spaceAfter=8,
        keepWithNext=1,
    )
    meta_style = ParagraphStyle(
        "KellaiMeta",
        parent=body,
        fontSize=8.8,
        leading=12,
        textColor=colors.HexColor("#4D5965"),
        alignment=TA_CENTER,
        spaceAfter=1.5,
        keepWithNext=1,
    )
    lead_style = ParagraphStyle(
        "KellaiLead",
        parent=body,
        fontSize=9.5,
        leading=14.2,
        backColor=colors.HexColor(f"#{LIGHT_BLUE}"),
        borderPadding=(7, 9, 7, 9),
        spaceBefore=7,
        spaceAfter=8,
    )
    notice_style = ParagraphStyle(
        "KellaiNotice",
        parent=body,
        fontName=heading_font,
        fontSize=8.8,
        leading=13.4,
        textColor=colors.HexColor(f"#{NAVY}"),
        backColor=colors.HexColor(f"#{LIGHT_GRAY}"),
        borderPadding=(6, 8, 6, 8),
        spaceBefore=4,
        spaceAfter=7,
    )
    heading1 = ParagraphStyle(
        "KellaiH1",
        parent=body,
        fontName=heading_font,
        fontSize=12.2,
        leading=17,
        textColor=colors.HexColor(f"#{NAVY}"),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=1,
    )
    heading2 = ParagraphStyle(
        "KellaiH2",
        parent=body,
        fontName=heading_font,
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=6,
        spaceAfter=3,
        keepWithNext=1,
    )
    numbered_style = ParagraphStyle(
        "KellaiNumbered",
        parent=body,
        leftIndent=18,
        firstLineIndent=-18,
        spaceAfter=4.2,
    )
    table_header = ParagraphStyle(
        "KellaiTableHeader",
        parent=body,
        fontName=heading_font,
        fontSize=8.4,
        leading=11.2,
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceAfter=0,
    )
    table_body = ParagraphStyle(
        "KellaiTableBody",
        parent=body,
        fontSize=7.8,
        leading=11.2,
        alignment=TA_LEFT,
        spaceAfter=0,
    )

    flowables = []
    counters: dict[int, int] = {}
    for block in iter_docx_blocks(source):
        if isinstance(block, DocxParagraph):
            raw = block.text.strip()
            if not raw:
                continue
            style_name = block.style.name if block.style is not None else "Normal"
            chosen = body
            if style_name == "Title":
                chosen = title_style
            elif style_name == "Legal Metadata":
                chosen = meta_style
            elif style_name == "Legal Lead":
                chosen = lead_style
            elif style_name == "Legal Notice":
                chosen = notice_style
            elif style_name == "Heading 1":
                chosen = heading1
            elif style_name == "Heading 2":
                chosen = heading2
            elif style_name == "Legal Numbered":
                chosen = numbered_style

            num_id = _docx_num_id(block)
            if num_id is not None:
                counters[num_id] = counters.get(num_id, 0) + 1
                raw = f"（{counters[num_id]}）{raw}"
            flowables.append(Paragraph(escape(raw), chosen))
        elif isinstance(block, DocxTable):
            data = []
            for row_idx, row in enumerate(block.rows):
                row_data = []
                for cell in row.cells:
                    text_value = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    row_data.append(
                        Paragraph(escape(text_value).replace("\n", "<br/>"), table_header if row_idx == 0 else table_body)
                    )
                data.append(row_data)
            if data:
                usable = LETTER[0] - 1.76 * inch
                if len(data[0]) == 3:
                    col_widths = [usable * 0.20, usable * 0.51, usable * 0.29]
                else:
                    col_widths = [usable / len(data[0])] * len(data[0])
                pdf_table = Table(
                    data,
                    colWidths=col_widths,
                    repeatRows=1,
                    hAlign="CENTER",
                    splitByRow=1,
                    splitInRow=0,
                )
                pdf_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
                            ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFBFC")]),
                            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{MID_GRAY}")),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                            ("TOPPADDING", (0, 0), (-1, -1), 5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                flowables.extend([pdf_table, Spacer(1, 5)])

    def on_page(canvas, pdf_doc):
        canvas.saveState()
        canvas.setTitle(source.core_properties.title or short_title)
        canvas.setAuthor(PROVIDER)
        canvas.setSubject(f"{PRODUCT}企业微信第三方应用法律文件")
        canvas.setFont(body_font, 7.7)
        canvas.setFillColor(colors.HexColor("#68727D"))
        y = LETTER[1] - 0.39 * inch
        canvas.drawString(0.88 * inch, y, f"{PRODUCT}  |  {PROVIDER}")
        right = f"{short_title}"
        canvas.drawRightString(LETTER[0] - 0.88 * inch, y, right)
        canvas.setStrokeColor(colors.HexColor(f"#{MID_GRAY}"))
        canvas.setLineWidth(0.35)
        canvas.line(0.88 * inch, y - 4, LETTER[0] - 0.88 * inch, y - 4)
        canvas.drawCentredString(LETTER[0] / 2, 0.34 * inch, f"第 {pdf_doc.page} 页")
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        rightMargin=0.88 * inch,
        leftMargin=0.88 * inch,
        topMargin=0.67 * inch,
        bottomMargin=0.58 * inch,
        title=source.core_properties.title,
        author=PROVIDER,
        subject=f"{PRODUCT}企业微信第三方应用法律文件",
        pageCompression=1,
    )
    pdf_doc.build(flowables, onFirstPage=on_page, onLaterPages=on_page)
    return pdf_path


def build_privacy_policy() -> Path:
    title = f"{PRODUCT}服务商隐私政策"
    doc, num_id = setup_document(title, "服务商隐私政策")
    add_title_block(
        doc,
        title,
        "本政策说明成都修茈科技有限公司作为客来来服务提供者，在企业通过企业微信安装、授权和使用客来来时如何处理个人信息与企业业务数据。请企业管理员在授权前完整阅读，并向本企业成员、外部联系人及其他相关个人履行必要的告知义务。",
    )
    add_notice(doc, "若您代表企业安装、授权或管理本应用，您确认已获得企业的合法授权，并已根据适用法律完成必要的内部审批、告知、同意或其他个人信息处理基础。")

    add_heading(doc, "一、适用范围与基本定义")
    add_para(doc, f"本政策适用于{PRODUCT}企业微信第三方应用、与其配套的桌面端、网页端、服务端接口及客户支持服务。企业自行部署的版本，以企业实际部署环境和双方书面约定为准。")
    add_para(doc, "“企业用户”指安装、授权或使用本服务的组织及其管理员；“个人用户”包括企业成员、外部联系人、客户及其他自然人；“企业业务数据”指企业在使用服务过程中导入、生成、接收或管理的数据，其中可能包含个人信息。")
    add_para(doc, "本政策不适用于企业微信平台自身、企业另行接入的模型服务商、支付机构或其他第三方独立提供的服务。相关第三方作为独立个人信息处理者时，其处理活动适用其自身规则。")

    add_heading(doc, "二、我们与企业用户的个人信息处理角色")
    add_para(doc, "对于企业账号注册、合同履行、计费结算、安全审计、客户支持等由我们自主决定处理目的和方式的活动，我们依法承担个人信息处理者责任。")
    add_para(doc, "对于企业在授权范围内同步、导入和管理的企业成员、外部联系人、客户、消息及业务记录，企业用户通常决定处理目的和方式，我们通常依据企业指令作为受托处理方处理。企业应确保其指令合法、正当、必要，并对向个人履行告知、取得同意或具备其他合法处理基础负责。")
    add_para(doc, "若特定场景中双方共同决定处理目的和方式，或法律法规对角色另有规定，双方将依法通过补充协议、产品提示或其他方式明确各自责任。")

    add_heading(doc, "三、我们处理的信息及使用目的")
    add_para(doc, "我们遵循目的明确、最小必要和权限最小化原则，仅在企业授权范围和实现相应功能所必要的期限内处理信息。具体范围会因企业开通功能、企业微信授权权限、部署方式及企业自主配置而不同。")
    add_info_table(
        doc,
        ("信息类别", "可能包含的具体字段", "主要用途"),
        (
            ("企业授权与应用身份信息", "企业 CorpID、企业名称、应用/Agent 标识、授权状态、授权时间、管理员或操作用户标识、安装回调状态；企业微信返回的授权信息。SuiteTicket、SuiteSecret、永久授权码等服务端凭据以加密或受控方式保存，不向普通客户端展示。", "完成第三方应用安装、换取访问凭证、维护授权关系、校验回调、同步授权范围内的数据。"),
            ("账号与组织信息", "手机号或电子邮箱、显示名称、头像、团队名称、成员角色、权限、登录和会话状态、邀请关系。", "注册登录、身份认证、团队协作、权限控制、账号安全和客户支持。"),
            ("企业成员与外部联系人信息", "企业成员 UserID、部门/角色等授权字段；外部联系人 external_userid、姓名、头像、类型、性别、UnionID（如平台返回）、跟进成员、备注、标签及关系状态。", "建立客户档案、识别联系人、分配跟进、维护客户关系和企业协作。"),
            ("消息、事件与沟通记录", "发送方/接收方标识、消息类型和内容、媒体文件引用、时间、送达或处理状态、回调事件、客户意向和跟进状态。", "统一收件箱、消息收发、客户服务、业务流程自动化、质检、审计和问题排查。"),
            ("客户与业务管理数据", "客户档案、联系人方式、备注、标签、线索阶段、跟进任务、订单或合作记录、导入文件、企业自行配置的字段。", "提供客户关系管理、销售协作、数据统计、提醒、导出和企业内部管理。"),
            ("AI 交互与模型调用数据", "用户输入、提示词、为完成任务所选取的客户或业务上下文、模型输出、模型名称、调用时间、状态、错误和用量信息。", "生成回复建议、意图分析、内容辅助、业务总结和服务质量优化。是否调用及调用哪一模型取决于企业功能选择和配置。"),
            ("设备、网络与安全日志", "IP 地址、设备/浏览器/操作系统类型、客户端版本、接口请求时间、操作记录、异常日志、鉴权和风控结果。", "保障系统稳定与安全、防止滥用、定位故障、审计关键操作并履行法定义务。"),
            ("合同、订单与服务支持信息", "企业联系人、联系方式、订单/版本、付款及发票状态、工单、沟通记录和反馈。", "订立和履行合同、计费结算、开票、交付、售后支持和争议处理。"),
        ),
        widths=(1.25, 3.45, 1.95),
    )

    add_heading(doc, "四、处理个人信息的合法性基础")
    add_para(doc, "我们根据具体场景，以取得个人同意、为订立或履行个人作为一方当事人的合同所必需、为依法制定的劳动规章制度和依法签订的集体合同实施人力资源管理所必需、履行法定义务、应对突发公共卫生事件或紧急情况下保护自然人生命健康和财产安全、在合理范围内处理个人自行公开或其他已合法公开的信息，以及法律法规规定的其他情形作为处理基础。")
    add_para(doc, "企业作为个人信息处理者向我们发出处理指令时，应自行确认相应处理基础。需要取得单独同意或书面同意的，企业和我们将根据各自角色依法完成。")

    add_heading(doc, "五、企业微信权限、设备权限与本地存储")
    add_numbered(
        doc,
        num_id,
        (
            "企业微信权限。我们仅在企业管理员选择并确认的权限范围内调用企业微信接口。企业可通过企业微信管理后台查看、调整或取消授权；取消某项权限可能导致对应功能不可用。",
            "设备权限。当用户主动使用扫码、上传附件、保存文件、消息提醒等功能时，客户端可能请求相机、相册/文件或通知权限。拒绝非必要权限不影响其他基础功能；用户可在系统设置中撤回。",
            "Cookie 与本地存储。网页端或桌面端可能使用会话 Cookie、Local Storage 或安全存储保存登录状态、界面偏好、设备标识和必要的防伪参数。我们不以此开展与服务无关的行为广告。",
        ),
    )

    add_heading(doc, "六、AI 功能与自动化处理")
    add_para(doc, "企业启用 AI 功能时，我们可能将完成所选任务所必要的输入、客户或业务上下文发送给企业选择或配置的模型服务。企业应在启用前评估模型服务商、部署区域和数据使用规则，并避免输入与任务无关的敏感个人信息、商业秘密或受特别限制的数据。")
    add_para(doc, "AI 输出具有概率性，可能不准确、不完整或不适合特定场景。涉及对外发送、客户承诺、定价、合同、财务、医疗、法律、人事决定或其他重要事项时，应由具备权限的人员复核。我们不会仅基于自动化决策对个人权益产生重大影响；如企业另行配置此类场景，应依法提供透明度、公平性和拒绝或说明机制。")
    add_para(doc, "依法需要标识生成合成内容时，我们将在产品能力范围内提供相应提示或技术支持，企业用户亦应在发布、传播或对外发送时履行适用的标识义务。")

    add_heading(doc, "七、委托处理、共享、转移与公开披露")
    add_para(doc, "我们不会出售个人信息。为提供服务，我们可能在最小必要范围内委托或向下列类型的接收方提供信息，并通过合同、权限控制、审计等措施约束其处理活动：")
    add_numbered(
        doc,
        num_id,
        (
            "企业微信及其运营主体：用于应用安装授权、身份校验、消息和事件回调、成员与外部联系人接口等平台能力。",
            "云计算、存储、网络、内容分发、短信、邮件、日志、安全与运维服务商：用于承载服务、发送验证信息、监测故障和防护攻击。",
            "企业选择或配置的 AI 模型服务商：仅在企业启用相关功能时，用于完成指定的生成、分析或总结任务。",
            "支付、开票、审计、法律和专业顾问：仅在订购、结算、合规、争议处理所必要的范围内。",
        ),
    )
    add_para(doc, "发生合并、分立、重组、资产转让、破产清算等导致个人信息转移的情形时，我们将依法告知接收方名称和联系方式，并要求接收方继续履行本政策；处理目的或方式发生变化的，接收方将依法重新取得同意。")
    add_para(doc, "除法律法规规定或为保护生命健康和财产安全所必需外，我们不会公开披露个人信息；确需公开披露时，将依法取得单独同意并进行显著告知。")

    add_heading(doc, "八、信息存储地点与跨境处理")
    add_para(doc, "本服务默认将个人信息存储在中华人民共和国境内，或存储于企业控制的本地设备、服务器和企业选择的境内部署环境。具体存储位置取决于产品版本和双方约定。")
    add_para(doc, "如企业主动选择境外模型、境外接口或其他可能造成个人信息出境的服务，我们会在可适用范围内提供必要提示，并与企业根据各自角色履行个人信息保护影响评估、认证、标准合同、申报安全评估或适用豁免、单独告知和单独同意等法定义务。在相关条件具备前，我们不会擅自为与服务无关的目的向境外提供个人信息。")

    add_heading(doc, "九、信息保存期限与删除")
    add_para(doc, "我们仅在实现本政策所述目的所必要的最短期限内保存信息，并综合考虑企业授权期限、合同和账号存续期、企业配置、业务连续性、安全审计、法定义务及争议处理需要。")
    add_numbered(
        doc,
        num_id,
        (
            "企业授权信息在授权有效期内保存；取消授权后，除为履行法定义务、结算、审计或争议处理确有必要外，原则上在三十日内删除或匿名化在线业务数据。",
            "账号与合同信息在账号或合同存续期内保存；账号注销或合同终止后，按适用法律和争议处理所需期限保存必要记录。",
            "网络运行与安全日志按照法律法规要求保存，通常不少于六个月。",
            "企业业务数据的具体保存期限可由企业配置或在双方协议中约定；备份中的残留副本将在备份轮换周期内安全清除，期间不再用于日常业务处理。",
        ),
    )
    add_para(doc, "如法律法规要求延长保存期限，或司法、行政机关依法要求冻结、保全，我们将按要求处理，并在目的达成后及时删除或匿名化。")

    add_heading(doc, "十、我们采取的安全措施")
    add_para(doc, "我们根据数据敏感程度和风险采取访问控制、最小权限、身份认证、传输加密、凭据加密或安全存储、关键操作日志、备份恢复、漏洞修复、员工保密和应急响应等措施。企业微信 SuiteSecret、SuiteTicket、永久授权码、访问令牌及企业配置的 API Key 等凭据仅用于服务端受控调用，并尽量避免在客户端、日志和支持沟通中明文展示。")
    add_para(doc, "互联网环境无法保证绝对安全。发生或可能发生个人信息泄露、篡改、丢失时，我们将立即采取补救措施，并依法向相关个人和主管部门告知事件类型、可能影响、已采取或拟采取的措施、个人可采取的防范措施及联系方式；法律允许不逐一告知时，依法公告。")

    add_heading(doc, "十一、个人信息主体的权利")
    add_para(doc, "个人依法享有知情、决定、限制或拒绝处理，查阅、复制、更正、补充、删除，撤回同意，注销账号，要求解释处理规则，以及在符合条件时请求转移个人信息等权利。")
    add_para(doc, "企业业务数据通常由企业控制。企业成员、外部联系人或客户可先向相关企业提出请求；企业可通过产品功能处理，或联系我们协助完成。对于由我们独立处理的账号、安全和服务支持信息，可直接发送邮件至本政策所列邮箱。为保护信息安全，我们可能核验身份和请求权限。")
    add_para(doc, "我们原则上在收到完整请求并完成身份核验后十五个工作日内答复。对于重复、明显无依据、需要过多技术手段、可能损害他人合法权益或法律法规允许拒绝的请求，我们可能说明理由后不予处理。")

    add_heading(doc, "十二、未成年人保护")
    add_para(doc, "本服务主要面向企业和具有完全民事行为能力的用户，不以不满十四周岁的未成年人为目标对象。企业不应在无合法处理基础和必要保护措施的情况下向本服务导入未成年人信息。若我们发现未经监护人同意处理了不满十四周岁未成年人的个人信息，将依法核实并尽快删除或采取其他保护措施。")

    add_heading(doc, "十三、企业用户的责任")
    add_numbered(
        doc,
        num_id,
        (
            "遵循合法、正当、必要和诚信原则，明确处理目的，限定授权人员和权限，避免超范围导入或使用个人信息。",
            "依法向企业成员、外部联系人、客户等履行告知义务，并在需要时取得同意、单独同意或书面同意。",
            "建立内部账号、数据导出、消息发送、AI 使用和离职交接制度，及时停用不再需要的账号和权限。",
            "不得通过本服务非法获取、买卖、泄露、骚扰、歧视或不当画像个人，不得将 AI 输出直接用于违法或高风险决定。",
        ),
    )

    add_heading(doc, "十四、政策更新")
    add_para(doc, "我们可能因法律法规、产品功能、处理目的或安全措施变化更新本政策。发生重大变化时，我们会通过企业微信授权页面、产品内通知、网站公告或其他显著方式告知；依法需要重新取得同意的，我们会在继续处理前完成。重大变化包括个人信息处理目的、方式、种类、保存期限、接收方或个人权利行使方式的实质变化。")

    add_heading(doc, "十五、联系我们")
    add_para(doc, f"个人信息保护负责人及服务商：{PROVIDER}。网站：{WEBSITE}。联系邮箱：{EMAIL}。邮件主题建议注明“客来来隐私权利请求”，并说明所涉企业、账号、请求类型和便于核验的信息。")
    add_para(doc, "如对我们的答复不满意，您可向有管辖权的网信、市场监管、公安等主管部门投诉举报，或依法通过诉讼等途径解决。")

    add_heading(doc, "十六、适用法律与规则")
    add_para(doc, "本政策依据并结合《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》《中华人民共和国网络安全法》《网络数据安全管理条例》《生成式人工智能服务管理暂行办法》《人工智能生成合成内容标识办法》及其他现行适用法律法规制定。如法律法规更新，以届时有效的强制性规定为准。")

    return save(doc, "客来来_服务商隐私政策_企业微信上传版.docx")


def build_service_agreement() -> Path:
    title = f"{PRODUCT}第三方服务协议"
    doc, num_id = setup_document(title, "第三方服务协议")
    add_title_block(
        doc,
        title,
        f"本协议由安装、授权或使用{PRODUCT}的企业用户（“甲方”）与{PROVIDER}（“乙方”）订立，用于明确企业微信第三方应用及配套服务的使用条件、数据责任、知识产权、费用、服务变更和责任边界。",
    )
    add_notice(doc, "甲方管理员点击安装、授权、确认、购买，或甲方实际使用本服务，即表示甲方已阅读、理解并同意本协议。管理员应确保有权代表甲方作出该等意思表示。限制或免除责任、数据处理、AI 使用、费用和争议解决条款已采用加粗、提示框或独立章节提示，请重点阅读。")

    add_heading(doc, "一、协议主体、生效与适用范围")
    add_para(doc, f"甲方为通过企业微信或乙方产品界面安装、授权、订购或使用{PRODUCT}的企业、组织或其他主体。乙方为{PROVIDER}。")
    add_para(doc, "本协议自甲方管理员完成安装或授权、点击同意、签署订单/合同，或甲方首次实际使用服务之最早发生时生效。甲方管理员的线上操作、企业微信授权记录、系统日志、订单和电子确认可作为协议订立与履行的证据。")
    add_para(doc, "如双方另行签署盖章合同、订单、数据处理协议或补充协议，文件之间不一致的，按补充协议、订单/盖章合同、本协议的顺序适用；但法律强制性规定另有要求的除外。")

    add_heading(doc, "二、定义")
    add_numbered(
        doc,
        num_id,
        (
            f"“本服务”指乙方通过{PRODUCT}企业微信第三方应用、桌面端、网页端、服务端接口或双方约定方式提供的客户管理、消息协同、跟进任务、数据统计、AI 辅助、系统集成和相关技术支持。",
            "“企业管理员”指甲方在企业微信或本服务中配置的超级管理员、应用管理员、团队所有者或具有相应权限的人员。",
            "“甲方数据”指甲方及其用户在使用服务时提交、导入、同步、生成或管理的数据，包括可能含有个人信息的企业成员、外部联系人、消息、客户和业务记录。",
            "“第三方服务”指企业微信平台、云服务、短信/邮件、支付、模型服务及其他由第三方独立提供并被本服务调用、链接或集成的服务。",
        ),
    )

    add_heading(doc, "三、服务内容与交付方式")
    add_para(doc, "乙方根据甲方开通的版本、授权权限、订单和实际可用能力提供服务。功能可能包括企业微信安装授权、联系人或客户同步、消息与事件回调、客户档案、线索和跟进管理、团队协作、内容与回复建议、AI 分析、数据统计、导入导出及技术支持。未在订单、产品页面或书面承诺中明确的功能不视为交付义务。")
    add_para(doc, "本服务可采用乙方托管、甲方本地部署或双方约定的其他方式。不同部署方式在数据存储位置、运维责任、备份、升级和可用性方面可能不同，以订单或交付文档为准。")
    add_para(doc, "乙方有权在不实质减损甲方已购买核心功能的前提下优化界面、技术架构和非核心功能。涉及重大不利变化的，乙方将提前以合理方式通知。")

    add_heading(doc, "四、企业账号、安装授权与管理员责任")
    add_numbered(
        doc,
        num_id,
        (
            "甲方应提供真实、准确、完整并持续有效的企业、管理员、联系人、开票和订单信息，不得冒用他人身份或无权代表的企业。",
            "甲方确认企业管理员有权选择应用权限、授权数据范围、配置第三方接口、邀请成员、分配角色、导出数据和终止服务。乙方可合理信赖经有效账号和企业微信授权发出的指令。",
            "甲方应妥善保管账号、验证码、登录凭据、API Key、企业微信密钥和设备，不得共享给无权人员。发现泄露、异常登录、离职未交接或权限错误时，应立即修改、撤回授权并通知乙方。",
            "甲方不得绕过权限控制、破解、扫描、攻击、干扰服务，不得利用本服务获取未授权数据或向第三方转售账号、授权能力和接口配额。",
        ),
    )

    add_heading(doc, "五、甲方数据和个人信息合规")
    add_para(doc, "甲方对甲方数据的来源、内容、处理目的和指令负责。甲方应确保其具有处理企业成员、外部联系人、客户和其他个人信息的合法性基础，并已完成必要的告知、同意、单独同意、书面同意、授权、影响评估或内部审批。")
    add_para(doc, "在甲方决定处理目的和方式的场景中，甲方通常为个人信息处理者，乙方按甲方指令受托处理。乙方仅在提供服务、履行协议、保障安全和法定义务所必要范围内处理甲方数据，不得擅自改变处理目的。乙方因账号安全、合同、结算、支持和合规而独立处理的信息，适用《客来来服务商隐私政策》。")
    add_numbered(
        doc,
        num_id,
        (
            "甲方不得导入来源不明、非法购买、超范围收集或无权处理的数据，不得利用本服务实施骚扰营销、诈骗、歧视、不当画像或其他侵害个人权益的行为。",
            "如甲方指令可能违反法律法规、监管要求、平台规则或侵害他人权益，乙方有权要求说明、暂停相关处理或拒绝执行，并及时通知甲方；紧急风险或法律禁止通知的除外。",
            "个人提出查阅、复制、更正、删除、撤回同意等请求时，甲方应依法处理；乙方在技术可行范围内提供合理协助。",
            "甲方应采取最小权限、账号回收、导出审批、员工培训和终端安全等措施，防止内部滥用和数据泄露。",
        ),
    )

    add_heading(doc, "六、AI 功能特别条款")
    add_notice(doc, "AI 功能提供辅助性建议，不替代甲方人员的专业判断。甲方应在对外发送、作出承诺或据此采取重要行动前进行人工复核。", label="AI 风险提示：")
    add_numbered(
        doc,
        num_id,
        (
            "企业启用 AI 功能时，乙方可能将完成指定任务所必要的输入、提示词和经选择的业务上下文提交给企业选择或配置的模型服务商。甲方应审慎选择模型、部署区域和数据使用设置，避免输入不必要的敏感个人信息、国家秘密、商业秘密或无权披露的信息。",
            "AI 输出可能包含事实错误、偏差、遗漏、不当内容或与甲方场景不适配的信息。甲方对是否采用、修改、发布或发送输出承担审核责任，不得将未经复核的输出直接用于医疗、法律、财务、信贷、人事、定价、合同承诺或其他可能对个人权益产生重大影响的决定。",
            "甲方发布或传播生成合成内容时，应依照适用法律和平台规则进行显式或隐式标识，不得删除、篡改或规避依法设置的标识。",
            "在法律允许范围内，乙方不保证 AI 输出绝对准确、完整、唯一、无偏见或不侵权；但乙方将提供必要的安全控制、投诉反馈和故障处置机制。",
        ),
    )

    add_heading(doc, "七、服务费用、订购与税费")
    add_para(doc, "服务价格、版本、席位、用量、服务期、付款节点、续费、退款和税费以产品订购页、报价单、订单或双方另行签署的合同为准。未明确标价的定制开发、数据迁移、驻场、培训和第三方用量费用不包含在基础服务费中。")
    add_para(doc, "甲方应按约定及时付款。逾期超过合理催告期限仍未支付的，乙方可暂停付费功能；暂停前将尽量通知并为甲方保留合理补救期。甲方仍应支付已发生费用及依法应承担的税费。")
    add_para(doc, "免费试用、测试、赠送或预览功能可在通知后调整或终止，除法律另有规定或乙方存在故意、重大过失外，乙方不对免费功能作持续可用性承诺。")

    add_heading(doc, "八、知识产权")
    add_para(doc, f"{PRODUCT}的软件、界面、商标、文档、算法、模型编排、模板、技术方案和其他成果的知识产权归乙方或相关权利人所有。甲方取得的是在服务期、授权范围和约定用途内的有限、非排他、不可转授权使用权，不因本协议取得底层软件或源代码所有权。")
    add_para(doc, "甲方保留其合法拥有的甲方数据、商标、素材和业务内容的权利。为提供服务，甲方授予乙方在服务期内、按甲方指令复制、传输、存储、格式转换和处理甲方数据的必要授权。除非另有书面约定，乙方不会以识别特定企业或个人的方式将甲方数据用于公开训练通用模型。")
    add_para(doc, "甲方向乙方提供建议、反馈或去标识化的功能需求时，乙方可将其用于改进产品，但不得据此披露甲方商业秘密或可识别个人的信息。")

    add_heading(doc, "九、第三方服务与企业微信平台")
    add_para(doc, "本服务依赖企业微信接口、网络、云基础设施及甲方选择的第三方服务。第三方可能调整接口、权限、配额、审核规则、价格或可用区域，或发生中断。乙方将尽合理努力适配、告知和恢复，但无法控制第三方的独立行为。")
    add_para(doc, "甲方使用第三方服务应同时遵守其协议、隐私规则和平台规范。因甲方未取得第三方账号、权限、配额或许可，第三方封禁甲方账号，或甲方自行配置错误导致的功能异常，不视为乙方违约；乙方应在合理范围内协助定位。")

    add_heading(doc, "十、服务运行、维护与变更")
    add_para(doc, "乙方将采取与服务风险相适应的技术和组织措施维护系统。因升级、迁移、安全处置或基础设施维护需要暂停服务时，乙方将尽量提前通知；紧急漏洞、攻击、监管要求或不可预见故障除外。")
    add_para(doc, "甲方应使用受支持的客户端版本并及时安装安全更新。因甲方拒绝升级、擅自修改部署环境、关闭安全组件或使用不兼容系统造成的风险和损失，由甲方承担相应责任。")

    add_heading(doc, "十一、保密义务")
    add_para(doc, "双方对在履行过程中获悉的对方商业秘密、未公开技术资料、客户资料、价格、账号凭据和其他明确标注或依性质应属保密的信息承担保密义务，仅向履约所必需且受不低于本协议保密义务约束的人员披露。")
    add_para(doc, "下列信息不属于保密信息：接收方无过错已公开的信息；接收方在披露前已合法知悉的信息；从无保密义务的第三方合法获得的信息；接收方独立开发且未使用披露方保密信息形成的信息。因法律、监管或司法要求必须披露的，接收方可在法定范围内披露，并在允许时提前通知披露方。")
    add_para(doc, "保密义务自信息披露之日起生效，在协议终止后持续五年；构成商业秘密或法律要求持续保护的信息，保密义务持续至其不再受保护或法律允许终止之日。")

    add_heading(doc, "十二、数据安全与事件响应")
    add_para(doc, "乙方将采取访问控制、身份认证、传输加密、服务端凭据保护、日志审计、备份恢复、漏洞管理和应急响应等措施。甲方应对其终端、管理员账号、内部网络、导出文件和本地部署环境采取相应保护。")
    add_para(doc, "任何一方发现可能影响甲方数据的泄露、篡改、丢失、越权访问或其他安全事件，应及时采取补救措施并通知另一方。双方将配合查明范围、保存证据、降低影响，并依据各自法定角色向个人和主管部门履行报告或告知义务。")

    add_heading(doc, "十三、暂停与终止")
    add_para(doc, "有下列情形之一的，乙方可根据风险程度限制、暂停或终止相关功能，并在法律允许和情况适当时提前或及时通知甲方：")
    add_numbered(
        doc,
        num_id,
        (
            "甲方严重违反法律法规、平台规则、本协议，或利用服务侵害他人权益；",
            "甲方账号、授权、接口或数据处理存在紧急安全风险，继续服务可能扩大损害；",
            "甲方逾期支付费用，经合理催告仍未补救；",
            "监管、司法机关、企业微信或其他有权主体依法要求暂停，或第三方基础能力终止导致服务无法继续；",
            "甲方提供虚假主体信息、无权代表企业授权，或以技术手段绕过安全与计费限制。",
        ),
    )
    add_para(doc, "甲方可通过企业微信管理后台取消授权，并按产品或订单规则申请终止。终止前，甲方应自行导出所需数据或联系乙方协助。终止后，乙方将根据隐私政策、双方约定和法律要求删除、匿名化或依法留存相关数据。")

    add_heading(doc, "十四、陈述与保证")
    add_para(doc, "双方均保证其依法设立并有权签署、履行本协议。乙方保证其将在适用法律和约定范围内提供服务，并对其人员和受托方采取合理管理措施。甲方保证其授权、数据来源、使用目的、内容和对外发送行为合法，不侵犯第三方合法权益。")

    add_heading(doc, "十五、免责声明与责任限制")
    add_notice(doc, "本章旨在合理分配可预见风险，不排除或限制因故意、重大过失造成人身损害，或法律明确禁止排除、限制的责任。", label="责任条款提示：")
    add_para(doc, "在法律允许范围内，乙方不对下列情形导致的服务中断、数据延迟或损失承担违约责任，但将提供合理协助：不可抗力；公共网络或运营商故障；企业微信或其他第三方独立服务中断、规则变化；甲方设备、网络、配置、账号泄露或违规操作；甲方拒绝升级或未采纳已通知的安全措施。")
    add_para(doc, "除乙方书面明确承诺外，本服务按现状和可用状态提供，乙方不保证服务完全无错误、永不中断，或 AI 输出、第三方数据和自动化结果适合甲方的特定商业目的。")
    add_para(doc, "在法律允许范围内，任何一方不对对方的间接损失、可得利益损失、商誉损失或第三方惩罚性索赔承担责任，但该损失系违约方故意或重大过失、侵犯知识产权、违反保密义务、违法处理个人信息或法律另有强制规定造成的除外。")
    add_para(doc, "在法律允许范围内，乙方因同一事件或一系列相关事件承担的累计赔偿责任，以该事件发生前十二个月内甲方就受影响服务向乙方实际支付的服务费总额为上限；如服务完全免费，则以人民币一千元为上限。该上限不适用于法律禁止限制责任的情形。")

    add_heading(doc, "十六、违约与赔偿")
    add_para(doc, "一方违反本协议，应在收到通知后及时停止违约、采取补救措施，并赔偿对方因此遭受的可证明直接损失。甲方因无权授权、非法数据来源、违法内容、骚扰营销、侵犯知识产权或未履行个人信息保护义务导致乙方遭受第三方索赔、行政处罚或合理维权费用的，应在其过错和责任范围内赔偿乙方。乙方因违法处理甲方数据或未履行约定安全义务造成损害的，依法承担相应责任。")

    add_heading(doc, "十七、通知")
    add_para(doc, "乙方可通过企业微信授权页面、产品内消息、弹窗、网站公告、订单联系人邮箱或手机号发送服务通知。涉及费用、重大功能变化、暂停终止、数据安全或协议重大修改的，乙方将采用与事项重要程度相适应的显著方式。")
    add_para(doc, "甲方应保持管理员和订单联系人信息有效。通知发送至甲方最近一次提供的联系方式或企业微信管理界面后，按通常可到达时间视为送达；法律对送达另有规定的，从其规定。")

    add_heading(doc, "十八、协议更新")
    add_para(doc, "乙方可因法律、监管、平台规则、产品能力或商业模式变化更新本协议。涉及甲方重大权益的不利变更，乙方将提前合理期限通知。甲方不同意的，可在变更生效前停止使用并按约定终止；甲方在生效后继续使用的，视为接受更新，但法律要求另行取得同意的除外。")

    add_heading(doc, "十九、适用法律与争议解决")
    add_para(doc, "本协议的订立、效力、履行、解释和争议解决适用中华人民共和国大陆地区法律。双方应先通过友好协商解决争议；自一方书面提出协商之日起三十日内未解决的，任一方可向乙方住所地有管辖权的人民法院提起诉讼。")

    add_heading(doc, "二十、其他")
    add_para(doc, "本协议任何条款被认定无效、违法或不可执行，不影响其他条款的效力；双方应以最接近原商业目的且合法有效的条款替代。乙方未立即行使权利不构成放弃。未经对方书面同意，任何一方不得转让本协议主要权利义务，但因合并、分立、重组或业务整体转让进行的承继除外，承继方应继续履约并依法处理数据。")
    add_para(doc, "本协议可通过电子方式订立。可靠的电子签名、平台确认、企业微信授权记录、订单、电子邮件和系统日志，与纸质签署文件具有依法可认可的证据效力。")

    add_heading(doc, "二十一、联系方式")
    add_para(doc, f"乙方：{PROVIDER}。产品：{PRODUCT}。网站：{WEBSITE}。联系邮箱：{EMAIL}。涉及合同、数据保护或投诉事项时，请在邮件中注明企业名称、管理员身份、订单或应用信息及具体诉求。")

    add_heading(doc, "二十二、适用法律与规则说明")
    add_para(doc, "本协议依据并结合《中华人民共和国民法典》《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》《中华人民共和国网络安全法》《网络数据安全管理条例》《生成式人工智能服务管理暂行办法》《人工智能生成合成内容标识办法》及企业微信平台现行规则制定。如相关规定更新，以届时有效的强制性规定为准。")

    return save(doc, "客来来_第三方服务协议_企业微信上传版.docx")


def main() -> None:
    privacy_docx = build_privacy_policy()
    service_docx = build_service_agreement()
    files = [
        privacy_docx,
        service_docx,
        export_pdf_from_docx(
            privacy_docx,
            PDF_OUT / "客来来_服务商隐私政策_企业微信上传版.pdf",
            "服务商隐私政策",
        ),
        export_pdf_from_docx(
            service_docx,
            PDF_OUT / "客来来_第三方服务协议_企业微信上传版.pdf",
            "第三方服务协议",
        ),
    ]
    for path in files:
        print(path)


if __name__ == "__main__":
    main()
