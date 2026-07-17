#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将客来来合同 Markdown 转为 Word (.docx) 与 PDF。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    Preformatted,
    KeepTogether,
    PageBreak,
    HRFlowable,
)

ROOT = Path("/Users/a4243342/Desktop/客来来/docs")
OUT = ROOT / "export"
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
FONT_NAME = "ArialUnicode"


def register_font() -> None:
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))


def parse_md(text: str):
    """粗解析 markdown：返回 block 列表。"""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.strip() == "---":
            blocks.append(("hr", None))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", parse_table(table_lines)))
            continue
        if re.match(r"^[-*] ", line) or re.match(r"^\d+\. ", line):
            items = []
            while i < len(lines) and (
                re.match(r"^[-*] ", lines[i])
                or re.match(r"^\d+\. ", lines[i])
                or (lines[i].startswith("  ") and items)
            ):
                items.append(lines[i].rstrip())
                i += 1
            blocks.append(("list", items))
            continue
        if line.startswith("> "):
            quote = []
            while i < len(lines) and lines[i].startswith(">"):
                quote.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            blocks.append(("quote", "\n".join(quote)))
            continue
        # paragraph (maybe multi-line until blank)
        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("|") and not lines[i].startswith(">") and lines[i].strip() != "---" and not re.match(r"^[-*] ", lines[i]) and not re.match(r"^\d+\. ", lines[i]):
            # stop if next looks like heading-ish already handled
            if lines[i].startswith("### ") or lines[i].startswith("## ") or lines[i].startswith("# "):
                break
            para.append(lines[i])
            i += 1
        blocks.append(("p", "\n".join(para)))
    return blocks


def parse_table(table_lines):
    rows = []
    for idx, raw in enumerate(table_lines):
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if idx == 1 and all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells):
            continue  # separator
        rows.append(cells)
    return rows


def md_inline_to_plain(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = s.replace("☐", "□").replace("☑", "☑")
    return s


def md_inline_to_rl(s: str) -> str:
    """Markdown inline -> ReportLab XML-ish."""
    s = (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", s)
    s = s.replace("\n", "<br/>")
    return s


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial Unicode MS"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = "Arial Unicode MS"
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
            st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_runs_md(paragraph, text: str, size=11, bold=False):
    """简单处理 **bold** 片段。"""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Arial Unicode MS"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def build_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    blocks = parse_md(text)
    doc = Document()
    set_doc_font(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for kind, data in blocks:
        if kind == "h1":
            p = doc.add_heading(level=1)
            add_runs_md(p, md_inline_to_plain(data), size=18, bold=True)
        elif kind == "h2":
            p = doc.add_heading(level=2)
            add_runs_md(p, md_inline_to_plain(data), size=14, bold=True)
        elif kind == "h3":
            p = doc.add_heading(level=3)
            add_runs_md(p, md_inline_to_plain(data), size=12, bold=True)
        elif kind == "hr":
            p = doc.add_paragraph("—" * 32)
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        elif kind == "quote":
            for line in data.split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                add_runs_md(p, line, size=10)
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif kind == "list":
            for item in data:
                m = re.match(r"^(\d+)\.\s+(.*)", item)
                if m:
                    p = doc.add_paragraph(style="List Number")
                    add_runs_md(p, m.group(2), size=11)
                elif item.startswith("- ") or item.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    add_runs_md(p, item[2:], size=11)
                else:
                    p = doc.add_paragraph()
                    add_runs_md(p, item.strip(), size=11)
        elif kind == "table":
            rows = data
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    val = row[c_idx] if c_idx < len(row) else ""
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs_md(p, md_inline_to_plain(val), size=9, bold=(r_idx == 0))
            doc.add_paragraph()
        elif kind == "p":
            for line in data.split("\n"):
                p = doc.add_paragraph()
                add_runs_md(p, line, size=11)

    doc.save(out_path)


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CN_Title",
            fontName=FONT_NAME,
            fontSize=16,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor=colors.HexColor("#111111"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_H1",
            fontName=FONT_NAME,
            fontSize=14,
            leading=20,
            spaceBefore=14,
            spaceAfter=8,
            textColor=colors.HexColor("#111111"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_H2",
            fontName=FONT_NAME,
            fontSize=12,
            leading=17,
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.HexColor("#222222"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_H3",
            fontName=FONT_NAME,
            fontSize=11,
            leading=15,
            spaceBefore=8,
            spaceAfter=4,
            textColor=colors.HexColor("#333333"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_Body",
            fontName=FONT_NAME,
            fontSize=9.5,
            leading=14,
            spaceAfter=4,
            alignment=TA_JUSTIFY,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_Quote",
            fontName=FONT_NAME,
            fontSize=9,
            leading=13,
            leftIndent=10,
            textColor=colors.HexColor("#555555"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_Li",
            fontName=FONT_NAME,
            fontSize=9.5,
            leading=14,
            leftIndent=12,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_Cell",
            fontName=FONT_NAME,
            fontSize=8,
            leading=11,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CN_CellHeader",
            fontName=FONT_NAME,
            fontSize=8,
            leading=11,
        )
    )
    return styles


def build_pdf(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    blocks = parse_md(text)
    styles = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title=md_path.stem,
        author="成都修茈科技有限公司",
    )
    story = []
    first_h1 = True

    for kind, data in blocks:
        if kind == "h1":
            style = styles["CN_Title"] if first_h1 else styles["CN_H1"]
            first_h1 = False
            story.append(Paragraph(md_inline_to_rl(data), style))
        elif kind == "h2":
            story.append(Paragraph(md_inline_to_rl(data), styles["CN_H2"]))
        elif kind == "h3":
            story.append(Paragraph(md_inline_to_rl(data), styles["CN_H3"]))
        elif kind == "hr":
            story.append(Spacer(1, 4))
            story.append(
                HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC"))
            )
            story.append(Spacer(1, 6))
        elif kind == "quote":
            story.append(Paragraph(md_inline_to_rl(data), styles["CN_Quote"]))
        elif kind == "list":
            for item in data:
                m = re.match(r"^(\d+)\.\s+(.*)", item)
                if m:
                    story.append(
                        Paragraph(f"{m.group(1)}. {md_inline_to_rl(m.group(2))}", styles["CN_Li"])
                    )
                elif item.startswith("- ") or item.startswith("* "):
                    story.append(
                        Paragraph(f"• {md_inline_to_rl(item[2:])}", styles["CN_Li"])
                    )
                else:
                    story.append(Paragraph(md_inline_to_rl(item.strip()), styles["CN_Li"]))
        elif kind == "table":
            rows = data
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            data_tbl = []
            for r_idx, row in enumerate(rows):
                cells = []
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    st = styles["CN_CellHeader"] if r_idx == 0 else styles["CN_Cell"]
                    cells.append(Paragraph(md_inline_to_rl(val), st))
                data_tbl.append(cells)
            avail = A4[0] - 3.2 * cm
            col_w = avail / cols
            t = Table(data_tbl, colWidths=[col_w] * cols, repeatRows=1)
            t.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), FONT_NAME),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F0F0")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 8))
        elif kind == "p":
            story.append(Paragraph(md_inline_to_rl(data), styles["CN_Body"]))

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont(FONT_NAME, 8)
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.drawString(1.6 * cm, 1.0 * cm, "成都修茈科技有限公司 · 客来来合作协议草稿 · 非正式法律文本")
        canvas.drawRightString(A4[0] - 1.6 * cm, 1.0 * cm, f"第 {doc_.page} 页")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    register_font()
    OUT.mkdir(parents=True, exist_ok=True)

    files = [
        "合同方案一_低现金股权版.md",
        "合同方案二_服务费加分成版.md",
        "合同方案三_纯服务费里程碑版.md",
        "合同三方案_选择与谈判要点.md",
    ]

    # 合并三份合同 + 选择要点 为一份总册（用户要「一份」）
    combined_md = OUT / "_combined_contracts.md"
    parts = []
    guide = ROOT / "合同三方案_选择与谈判要点.md"
    parts.append(guide.read_text(encoding="utf-8"))
    parts.append("\n\n---\n\n")
    for name in files[:3]:
        parts.append(f"\n\n")
        parts.append((ROOT / name).read_text(encoding="utf-8"))
        parts.append("\n\n---\n\n")
    combined_md.write_text("".join(parts), encoding="utf-8")

    # 总册
    build_docx(combined_md, OUT / "客来来_合作协议三方案_修茈.docx")
    build_pdf(combined_md, OUT / "客来来_合作协议三方案_修茈.pdf")

    # 分册也各出一份，方便单独发
    for name in files:
        stem = Path(name).stem
        md = ROOT / name
        build_docx(md, OUT / f"{stem}.docx")
        build_pdf(md, OUT / f"{stem}.pdf")

    print("OK")
    for p in sorted(OUT.glob("*")):
        if p.name.startswith("_"):
            continue
        print(f"{p.name}\t{p.stat().st_size}")


if __name__ == "__main__":
    main()
